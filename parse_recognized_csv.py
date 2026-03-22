"""Parse OCR output CSVs into readable per-image text files.

Supports two formats produced by this repo:

1) TrOCR pipeline output (ocr_pipeline.py):
   - CSV columns: image_stem,index,score,crop_file,text
   - Produces one joined transcription per image_stem by sorting on index.

2) Qwen pipeline output (qwen_pipeline.py):
   - CSV columns: image_name,transcription
   - Produces one transcription per image_name (stem).

Writes:
- <output_dir>/<image_stem>.txt for each image
- <output_dir>/all_text.txt concatenating all images
- <output_dir>/manifest.csv with image key + output path
Optional:
- <output_dir>/transcriptions.docx (or custom path)

Example:
  python parse_recognized_csv.py \
    --input_csv Results_RealWorld_TrOCR_full/recognized_words.csv \
    --output_dir Parsed_TrOCR_full

  python parse_recognized_csv.py \
    --input_csv Results_RealWorld_Qwen/recognized_text.csv \
    --output_dir Parsed_Qwen
"""

from __future__ import annotations

import argparse
import csv
from pathlib import Path


def str2bool(v: object) -> bool:
    if isinstance(v, bool):
        return v
    return str(v).strip().lower() in {"1", "true", "t", "yes", "y"}


def _normalize_key(name_or_stem: str) -> str:
    # If passed a filename, drop extension; if passed a stem, return as-is.
    p = Path(name_or_stem)
    return p.stem if p.suffix else name_or_stem


def parse_trocr_recognized_words(rows: list[dict[str, str]]) -> dict[str, str]:
    """Parse rows from TrOCR recognized_words.csv into {image_stem: joined_text}."""
    items: list[tuple[str, int, str]] = []
    for r in rows:
        stem = (r.get("image_stem") or "").strip()
        if not stem:
            continue
        idx_str = (r.get("index") or "0").strip()
        try:
            idx = int(idx_str)
        except ValueError:
            idx = 0
        text = (r.get("text") or "").strip()
        items.append((stem, idx, text))

    items.sort(key=lambda t: (t[0], t[1]))

    out: dict[str, list[str]] = {}
    for stem, _idx, text in items:
        if stem not in out:
            out[stem] = []
        if text:
            out[stem].append(text)

    return {stem: " ".join(parts).strip() for stem, parts in out.items()}


def parse_qwen_recognized_text(rows: list[dict[str, str]]) -> dict[str, str]:
    """Parse rows from Qwen recognized_text.csv into {image_stem: transcription}."""
    out: dict[str, str] = {}
    for r in rows:
        name = (r.get("image_name") or "").strip()
        if not name:
            continue
        key = _normalize_key(name)
        transcription = (r.get("transcription") or "").strip()
        out[key] = transcription
    return out


def read_csv_rows(csv_path: Path) -> tuple[list[str], list[dict[str, str]]]:
    with csv_path.open(newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        rows = [dict(r) for r in reader]
    return fieldnames, rows


def write_outputs(text_by_image: dict[str, str], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    # Stable ordering
    keys = sorted(text_by_image.keys())

    manifest_path = output_dir / "manifest.csv"
    all_text_path = output_dir / "all_text.txt"

    with manifest_path.open("w", newline="", encoding="utf-8") as mf, all_text_path.open(
        "w", encoding="utf-8"
    ) as af:
        mw = csv.DictWriter(mf, fieldnames=["image_key", "output_txt", "num_chars"])
        mw.writeheader()

        for key in keys:
            text = (text_by_image.get(key) or "").strip()
            out_txt = output_dir / f"{key}.txt"
            out_txt.write_text(text + ("\n" if text and not text.endswith("\n") else ""), encoding="utf-8")

            # all_text.txt with separators
            af.write(f"===== {key} =====\n")
            af.write(text)
            af.write("\n\n")

            mw.writerow({"image_key": key, "output_txt": str(out_txt), "num_chars": len(text)})


def write_docx(text_by_image: dict[str, str], docx_path: Path, title: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise SystemExit(
            "Missing dependency: python-docx. Install with: python -m pip install python-docx"
        ) from e

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for key in sorted(text_by_image.keys()):
        text = (text_by_image.get(key) or "").strip()
        doc.add_heading(str(key), level=2)

        if text:
            # Preserve line breaks from the transcription.
            for line in text.splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

        doc.add_paragraph("_" * 60)

    doc.save(str(docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse OCR CSVs into per-image .txt outputs")
    parser.add_argument("--input_csv", required=True, help="Path to recognized_words.csv or recognized_text.csv")
    parser.add_argument("--output_dir", required=True, help="Folder to write per-image text files")
    parser.add_argument(
        "--write_docx",
        type=str2bool,
        default=False,
        help="If True, also write a DOCX file with one section per image.",
    )
    parser.add_argument(
        "--docx_path",
        default="",
        help="Optional output DOCX path. Default: <output_dir>/transcriptions.docx",
    )
    parser.add_argument(
        "--docx_title",
        default="Handwriting Recognition Results",
        help="Title shown at top of DOCX.",
    )

    args = parser.parse_args()

    csv_path = Path(args.input_csv)
    if not csv_path.exists():
        raise SystemExit(f"Input CSV not found: {csv_path}")

    output_dir = Path(args.output_dir)

    fieldnames, rows = read_csv_rows(csv_path)
    fieldset = {f.strip() for f in fieldnames if f}

    if {"image_stem", "index", "text"}.issubset(fieldset):
        text_by_image = parse_trocr_recognized_words(rows)
        mode = "trocr_recognized_words"
    elif {"image_name", "transcription"}.issubset(fieldset):
        text_by_image = parse_qwen_recognized_text(rows)
        mode = "qwen_recognized_text"
    else:
        raise SystemExit(
            "Unrecognized CSV format. Expected columns either: "
            "(image_stem,index,text,...) or (image_name,transcription). "
            f"Got columns: {fieldnames}"
        )

    write_outputs(text_by_image, output_dir)
    if args.write_docx:
        docx_path = Path(args.docx_path) if args.docx_path else (output_dir / "transcriptions.docx")
        write_docx(text_by_image, docx_path=docx_path, title=str(args.docx_title))
        print(f"Wrote DOCX: {docx_path}")
    print(f"Parsed {len(text_by_image)} images ({mode})")
    print(f"Wrote: {output_dir}")


if __name__ == "__main__":
    main()
